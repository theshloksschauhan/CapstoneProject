from dotenv import load_dotenv
from pathlib import Path
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

import os
import io
import logging
from datetime import datetime, timezone, timedelta
from typing import List, Optional

from fastapi import FastAPI, APIRouter, HTTPException, Request, Response, Depends, UploadFile, File, Query
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from starlette.middleware.cors import CORSMiddleware
from starlette.middleware.gzip import GZipMiddleware
from starlette.middleware.trustedhost import TrustedHostMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, EmailStr, Field
from bson import ObjectId
from pypdf import PdfReader
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

import auth
import llm_agents

# ---------------- DB ----------------
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

app = FastAPI(title="CareerOS API")
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

api_router = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

MAX_TEXT_LENGTH = 50_000
DAILY_AI_RUN_LIMIT = int(os.environ.get("DAILY_AI_RUN_LIMIT", "20"))


def _allowed_origins() -> list[str]:
    raw = os.environ.get("FRONTEND_URL", "http://localhost:3000")
    return [origin.strip() for origin in raw.split(",") if origin.strip()]


# ---------------- Models ----------------
class RegisterInput(BaseModel):
    name: str
    email: EmailStr
    password: str = Field(min_length=6)


class LoginInput(BaseModel):
    email: EmailStr
    password: str


class ProfileUpdate(BaseModel):
    name: Optional[str] = None
    current_password: Optional[str] = None
    new_password: Optional[str] = Field(default=None, min_length=6)


class ApplicationCreate(BaseModel):
    job_title: str = Field(max_length=500)
    company: str = Field(max_length=500)
    job_description: str = Field(max_length=MAX_TEXT_LENGTH)
    resume_text: str = Field(max_length=MAX_TEXT_LENGTH)


class StatusUpdate(BaseModel):
    status: str


VALID_STATUSES = ["draft", "applied", "interviewing", "offer", "rejected"]


# ---------------- Dependencies ----------------
async def current_user(request: Request) -> dict:
    return await auth.get_current_user_from_db(request, db)


async def admin_user(request: Request) -> dict:
    return await auth.require_admin(request, db)


def serialize_app(doc: dict) -> dict:
    doc["id"] = str(doc.pop("_id"))
    return doc


def _trim_app_for_list(d: dict) -> dict:
    d = serialize_app(d)
    d.pop("job_description", None)
    d.pop("resume_text", None)
    results = d.get("results", {})
    d["match_rate"] = results.get("fit_analysis", {}).get("match_rate")
    d["ats_overall"] = results.get("ats_score", {}).get("overall_score")
    d.pop("results", None)
    return d


async def _check_daily_ai_quota(user_id: str):
    today = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
    count = await db.applications.count_documents({
        "user_id": user_id,
        "generated": True,
        "updated_at": {"$gte": today.isoformat()},
    })
    if count >= DAILY_AI_RUN_LIMIT:
        raise HTTPException(
            status_code=429,
            detail=f"Daily AI generation limit reached ({DAILY_AI_RUN_LIMIT}). Try again tomorrow.",
        )


# ---------------- Auth Routes ----------------
@api_router.post("/auth/register")
async def register(payload: RegisterInput, response: Response):
    email = payload.email.lower()
    if await db.users.find_one({"email": email}):
        raise HTTPException(status_code=400, detail="Email already registered")
    user_doc = {
        "name": payload.name,
        "email": email,
        "password_hash": auth.hash_password(payload.password),
        "role": "user",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    result = await db.users.insert_one(user_doc)
    uid = str(result.inserted_id)
    access = auth.create_access_token(uid, email)
    refresh = auth.create_refresh_token(uid)
    auth.set_auth_cookies(response, access, refresh)
    return {"id": uid, "name": payload.name, "email": email, "role": "user", "access_token": access}


@api_router.post("/auth/login")
async def login(payload: LoginInput, request: Request, response: Response):
    email = payload.email.lower()
    ip = request.client.host if request.client else "unknown"
    identifier = f"{ip}:{email}"
    await auth.check_lockout(db, identifier)
    user = await db.users.find_one({"email": email})
    if not user or not auth.verify_password(payload.password, user["password_hash"]):
        await auth.register_failed_attempt(db, identifier)
        raise HTTPException(status_code=401, detail="Invalid email or password")
    await auth.clear_attempts(db, identifier)
    uid = str(user["_id"])
    access = auth.create_access_token(uid, email)
    refresh = auth.create_refresh_token(uid)
    auth.set_auth_cookies(response, access, refresh)
    return {"id": uid, "name": user.get("name"), "email": email,
            "role": user.get("role", "user"), "access_token": access}


@api_router.post("/auth/logout")
async def logout(response: Response, user: dict = Depends(current_user)):
    auth.clear_auth_cookies(response)
    return {"message": "Logged out"}


@api_router.get("/auth/me")
async def me(user: dict = Depends(current_user)):
    return user


@api_router.patch("/auth/profile")
async def update_profile(payload: ProfileUpdate, user: dict = Depends(current_user)):
    updates = {}
    if payload.name is not None:
        if not payload.name.strip():
            raise HTTPException(status_code=400, detail="Name cannot be empty")
        updates["name"] = payload.name.strip()

    if payload.new_password:
        if not payload.current_password:
            raise HTTPException(status_code=400, detail="Current password required")
        doc = await db.users.find_one({"_id": ObjectId(user["id"])})
        if not doc or not auth.verify_password(payload.current_password, doc["password_hash"]):
            raise HTTPException(status_code=400, detail="Current password is incorrect")
        updates["password_hash"] = auth.hash_password(payload.new_password)

    if not updates:
        raise HTTPException(status_code=400, detail="No changes provided")

    await db.users.update_one({"_id": ObjectId(user["id"])}, {"$set": updates})
    updated = await db.users.find_one({"_id": ObjectId(user["id"])})
    updated["id"] = str(updated.pop("_id"))
    updated.pop("password_hash", None)
    return updated


@api_router.post("/auth/refresh")
async def refresh_token(request: Request, response: Response):
    token = request.cookies.get("refresh_token")
    if not token:
        raise HTTPException(status_code=401, detail="No refresh token")
    import jwt as _jwt
    try:
        payload = _jwt.decode(token, auth.get_jwt_secret(), algorithms=[auth.JWT_ALGORITHM])
        if payload.get("type") != "refresh":
            raise HTTPException(status_code=401, detail="Invalid token type")
        user = await db.users.find_one({"_id": ObjectId(payload["sub"])})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        access = auth.create_access_token(str(user["_id"]), user["email"])
        auth.set_access_cookie(response, access)
        return {"access_token": access}
    except _jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid refresh token")


# ---------------- Health ----------------
@api_router.get("/health")
async def health():
    db_ok = False
    try:
        await db.command("ping")
        db_ok = True
    except Exception:
        pass
    return {
        "status": "ok" if db_ok else "degraded",
        "database": db_ok,
        "openai_configured": bool(os.environ.get("OPENAI_API_KEY") or os.environ.get("GEMINI_API_KEY") or os.environ.get("GROQ_API_KEY") or os.environ.get("OPENROUTER_API_KEY")),
    }


# ---------------- Resume Parsing ----------------
@api_router.post("/resume/parse")
async def parse_resume(file: UploadFile = File(...), user: dict = Depends(current_user)):
    content = await file.read()
    text = ""
    name = (file.filename or "").lower()
    try:
        if name.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(content))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        else:
            text = content.decode("utf-8", errors="ignore")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not parse file: {e}")
    text = text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="No readable text found in the file.")
    if len(text) > MAX_TEXT_LENGTH:
        raise HTTPException(status_code=400, detail=f"Resume exceeds {MAX_TEXT_LENGTH} character limit.")
    return {"filename": file.filename, "text": text, "char_count": len(text)}


# ---------------- Applications ----------------
@api_router.post("/applications")
async def create_application(payload: ApplicationCreate, user: dict = Depends(current_user)):
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "user_id": user["id"],
        "job_title": payload.job_title,
        "company": payload.company,
        "job_description": payload.job_description,
        "resume_text": payload.resume_text,
        "status": "draft",
        "results": {},
        "generated": False,
        "created_at": now,
        "updated_at": now,
        "status_history": [{"status": "draft", "at": now}],
    }
    result = await db.applications.insert_one(doc)
    doc["_id"] = result.inserted_id
    return serialize_app(doc)


@api_router.get("/applications")
async def list_applications(
    user: dict = Depends(current_user),
    page: int = Query(1, ge=1),
    limit: int = Query(50, ge=1, le=100),
):
    skip = (page - 1) * limit
    query = {"user_id": user["id"]}
    total = await db.applications.count_documents(query)
    cursor = db.applications.find(query).sort("created_at", -1).skip(skip).limit(limit)
    docs = await cursor.to_list(limit)
    return {
        "items": [_trim_app_for_list(d) for d in docs],
        "total": total,
        "page": page,
        "limit": limit,
    }


@api_router.get("/applications/{app_id}")
async def get_application(app_id: str, user: dict = Depends(current_user)):
    try:
        doc = await db.applications.find_one({"_id": ObjectId(app_id), "user_id": user["id"]})
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid application id")
    if not doc:
        raise HTTPException(status_code=404, detail="Application not found")
    return serialize_app(doc)


@api_router.patch("/applications/{app_id}/status")
async def update_status(app_id: str, payload: StatusUpdate, user: dict = Depends(current_user)):
    if payload.status not in VALID_STATUSES:
        raise HTTPException(status_code=400, detail="Invalid status")
    now = datetime.now(timezone.utc).isoformat()
    result = await db.applications.update_one(
        {"_id": ObjectId(app_id), "user_id": user["id"]},
        {"$set": {"status": payload.status, "updated_at": now},
         "$push": {"status_history": {"status": payload.status, "at": now}}},
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Application not found")
    doc = await db.applications.find_one({"_id": ObjectId(app_id)})
    return serialize_app(doc)


@api_router.delete("/applications/{app_id}")
async def delete_application(app_id: str, user: dict = Depends(current_user)):
    result = await db.applications.delete_one({"_id": ObjectId(app_id), "user_id": user["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Application not found")
    return {"message": "Deleted"}


async def _fetch_owned_app(app_id: str, user: dict) -> dict:
    try:
        doc = await db.applications.find_one({"_id": ObjectId(app_id), "user_id": user["id"]})
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid application id")
    if not doc:
        raise HTTPException(status_code=404, detail="Application not found")
    return doc


@api_router.post("/applications/{app_id}/run")
@limiter.limit("10/hour")
async def run_pipeline(request: Request, app_id: str, user: dict = Depends(current_user)):
    await _check_daily_ai_quota(user["id"])
    doc = await _fetch_owned_app(app_id, user)
    try:
        results = await llm_agents.run_full_pipeline(
            doc["resume_text"], doc["job_description"], doc["company"], doc["job_title"], app_id
        )
    except Exception as e:
        logger.exception("Pipeline failed")
        raise HTTPException(status_code=500, detail=f"AI generation failed: {e}")
    now = datetime.now(timezone.utc).isoformat()
    await db.applications.update_one(
        {"_id": ObjectId(app_id)},
        {"$set": {"results": results, "generated": True, "updated_at": now}},
    )
    updated = await db.applications.find_one({"_id": ObjectId(app_id)})
    return serialize_app(updated)


@api_router.post("/applications/{app_id}/regenerate/{agent}")
@limiter.limit("20/hour")
async def regenerate_agent(request: Request, app_id: str, agent: str, user: dict = Depends(current_user)):
    if agent not in llm_agents.AGENT_MAP:
        raise HTTPException(status_code=400, detail="Unknown agent")
    await _check_daily_ai_quota(user["id"])
    doc = await _fetch_owned_app(app_id, user)
    fit_context = doc.get("results", {}).get("fit_analysis")
    try:
        result = await llm_agents.AGENT_MAP[agent](
            doc["resume_text"], doc["job_description"], doc["company"], doc["job_title"], app_id,
            fit_context=fit_context,
        )
    except Exception as e:
        logger.exception("Regeneration failed")
        raise HTTPException(status_code=500, detail=f"AI generation failed: {e}")
    now = datetime.now(timezone.utc).isoformat()
    await db.applications.update_one(
        {"_id": ObjectId(app_id)},
        {"$set": {f"results.{agent}": result, "generated": True, "updated_at": now}},
    )
    updated = await db.applications.find_one({"_id": ObjectId(app_id)})
    return serialize_app(updated)


def _docx_from_cover_letter(app_doc: dict) -> bytes:
    cover = app_doc.get("results", {}).get("cover_letter", {}).get("cover_letter", "").strip()
    if not cover:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    document = Document()
    document.add_heading(f"{app_doc.get('job_title', 'Cover Letter')} at {app_doc.get('company', '')}".strip(), level=1)
    for paragraph in cover.splitlines():
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())
        else:
            document.add_paragraph("")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_from_resume(app_doc: dict) -> bytes:
    resume = app_doc.get("results", {}).get("resume_rewrite", {}).get("rewritten_resume", "").strip()
    if not resume:
        raise HTTPException(status_code=404, detail="Rewritten resume not found")
    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=18,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=6,
    )
    story = [
        Paragraph(f"{app_doc.get('job_title', 'Resume')} at {app_doc.get('company', '')}".strip(), title_style),
        Spacer(1, 0.12 * inch),
    ]
    for paragraph in resume.splitlines():
        if paragraph.strip():
            safe = paragraph.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, body_style))
        else:
            story.append(Spacer(1, 0.12 * inch))
    pdf.build(story)
    return buffer.getvalue()


@api_router.get("/applications/{app_id}/export/cover-letter.docx")
async def export_cover_letter(app_id: str, user: dict = Depends(current_user)):
    doc = await _fetch_owned_app(app_id, user)
    content = _docx_from_cover_letter(doc)
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="cover-letter-{app_id}.docx"'},
    )


@api_router.get("/applications/{app_id}/export/resume.pdf")
async def export_resume(app_id: str, user: dict = Depends(current_user)):
    doc = await _fetch_owned_app(app_id, user)
    content = _pdf_from_resume(doc)
    return Response(
        content=content,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="resume-{app_id}.pdf"'},
    )


@api_router.get("/applications/{app_id}/export/followup.ics")
async def export_followup_ics(app_id: str, user: dict = Depends(current_user)):
    doc = await _fetch_owned_app(app_id, user)
    job_title = doc.get("job_title", "Job")
    company = doc.get("company", "Company")
    # Follow up 1 week after creation
    try:
        created_at = datetime.fromisoformat(doc["created_at"])
    except:
        created_at = datetime.now(timezone.utc)
    
    start_time = created_at + timedelta(days=7)
    end_time = start_time + timedelta(minutes=30)
    
    dtstamp = datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')
    dtstart = start_time.strftime('%Y%m%dT%H%M%SZ')
    dtend = end_time.strftime('%Y%m%dT%H%M%SZ')

    ics_content = (
        "BEGIN:VCALENDAR\n"
        "VERSION:2.0\n"
        "PRODID:-//CareerOS//Followup//EN\n"
        "BEGIN:VEVENT\n"
        f"UID:followup-{app_id}@careeros.local\n"
        f"DTSTAMP:{dtstamp}\n"
        f"DTSTART:{dtstart}\n"
        f"DTEND:{dtend}\n"
        f"SUMMARY:Follow up on application: {job_title} at {company}\n"
        f"DESCRIPTION:It has been 1 week since you applied for {job_title} at {company}. Time to send a polite follow-up email!\n"
        "END:VEVENT\n"
        "END:VCALENDAR"
    )

    return Response(
        content=ics_content,
        media_type="text/calendar",
        headers={"Content-Disposition": f'attachment; filename="followup-{app_id}.ics"'},
    )


@api_router.post("/jd/scrape")
async def scrape_jd(payload: dict, user: dict = Depends(current_user)):
    url = (payload.get("url") or "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="URL is required")
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="URL must start with http:// or https://")
    try:
        import requests
        from bs4 import BeautifulSoup

        response = requests.get(
            url,
            timeout=10,
            headers={"User-Agent": "CareerOS/1.0 (+https://careeros.local)"},
            stream=True,
        )
        content_type = response.headers.get("content-type", "")
        if "text/html" not in content_type and "application/xhtml+xml" not in content_type:
            raise HTTPException(status_code=400, detail="URL did not return HTML content")
        raw = response.content[:1_000_000]
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = "\n".join(line.strip() for line in soup.get_text("\n").splitlines() if line.strip())
        if not text:
            raise HTTPException(status_code=400, detail="Could not extract readable text from the page")
        return {"text": text[:MAX_TEXT_LENGTH], "char_count": min(len(text), MAX_TEXT_LENGTH)}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not scrape URL: {exc}")


@api_router.get("/stats")
async def stats(user: dict = Depends(current_user)):
    docs = await db.applications.find({"user_id": user["id"]}).to_list(500)
    total = len(docs)
    by_status = {}
    match_rates = []
    ats_scores = []
    for d in docs:
        by_status[d.get("status", "draft")] = by_status.get(d.get("status", "draft"), 0) + 1
        results = d.get("results", {})
        mr = results.get("fit_analysis", {}).get("match_rate")
        ats = results.get("ats_score", {}).get("overall_score")
        if isinstance(mr, (int, float)):
            match_rates.append(mr)
        if isinstance(ats, (int, float)):
            ats_scores.append(ats)
    return {
        "total": total,
        "by_status": by_status,
        "avg_match_rate": round(sum(match_rates) / len(match_rates)) if match_rates else None,
        "avg_ats": round(sum(ats_scores) / len(ats_scores)) if ats_scores else None,
        "generated_count": sum(1 for d in docs if d.get("generated")),
    }


@api_router.get("/admin/stats")
async def admin_stats(admin: dict = Depends(admin_user)):
    user_count = await db.users.count_documents({})
    app_count = await db.applications.count_documents({})
    generated_count = await db.applications.count_documents({"generated": True})
    return {
        "users": user_count,
        "applications": app_count,
        "generated_applications": generated_count,
    }


@api_router.get("/")
async def root():
    return {"message": "CareerOS API running"}


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_origins=_allowed_origins(),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.add_middleware(GZipMiddleware, minimum_size=1000)

@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    return response

@app.on_event("startup")
async def startup():
    await db.users.create_index("email", unique=True)
    await db.login_attempts.create_index("identifier")
    await auth.seed_admin(db)
    logger.info("Startup complete")


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
